import io
import logging
import re
from pypdf import PdfReader

logger = logging.getLogger("ResumeParser")

def clean_extracted_text(text: str) -> str:
    """
    Cleans and normalizes extracted text to optimize LLM token usage:
    - Replaces 3 or more consecutive newlines with 2 newlines.
    - Replaces multiple whitespace/tabs with a single space.
    - Strips leading/trailing whitespaces.
    
    :param text: Raw extracted string from document
    :return: Normalized, token-efficient text string
    """
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    return text.strip()


def extract_text_from_pdf(file_bytes: bytes, max_chars: int = 50000) -> tuple[str, int]:
    """
    Extracts text content and page metadata from a PDF file provided as bytes.
    Limits total character ingestion to max_chars to avoid excessive token overhead.

    :param file_bytes: Raw binary payload of the PDF document
    :param max_chars: Maximum character limit for ingestion (default: 50,000)
    :return: Tuple of (cleaned_extracted_text, total_page_count)
    :raises ValueError: When no extractable text is found (e.g., scanned/image-only PDF)
    """
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        total_pages = len(reader.pages)
        
        extracted_text = []
        char_count = 0
        
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if char_count + len(text) > max_chars:
                # Add partial text to fit limit
                text_slice = text[:max_chars - char_count]
                extracted_text.append(text_slice)
                logger.info(f"PDF text extraction truncated at page {i+1} due to character limit.")
                break
            
            extracted_text.append(text)
            char_count += len(text)
            
        full_text = clean_extracted_text("\n".join(extracted_text))
        
        if not full_text:
            raise ValueError("No text could be extracted from the PDF. The file might be scanned or image-only.")
            
        return full_text, total_pages
        
    except Exception as e:
        logger.error(f"Error extracting PDF: {str(e)}")
        raise e

def extract_text_from_docx(file_bytes: bytes, max_chars: int = 50000) -> str:
    """
    Extracts text content from a DOCX file provided as bytes.
    """
    try:
        import docx
        doc_file = io.BytesIO(file_bytes)
        doc = docx.Document(doc_file)
        
        extracted_text = []
        char_count = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text or ""
            if char_count + len(text) + 1 > max_chars:
                text_slice = text[:max_chars - char_count]
                extracted_text.append(text_slice)
                break
            extracted_text.append(text)
            char_count += len(text) + 1
            
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                text = " | ".join(row_text)
                if char_count + len(text) + 1 > max_chars:
                    text_slice = text[:max_chars - char_count]
                    extracted_text.append(text_slice)
                    break
                extracted_text.append(text)
                char_count += len(text) + 1
                
        full_text = clean_extracted_text("\n".join(extracted_text))
        
        if not full_text:
            raise ValueError("No text could be extracted from the DOCX file.")
            
        return full_text
    except Exception as e:
        logger.error(f"Error extracting DOCX: {str(e)}")
        raise e

def extract_text_from_txt(file_bytes: bytes, max_chars: int = 50000) -> str:
    """
    Extracts text content from a TXT file provided as bytes.
    """
    try:
        # Try decoding as utf-8, fallback to latin-1
        try:
            text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            text = file_bytes.decode('latin-1')
            
        full_text = clean_extracted_text(text[:max_chars])
        
        if not full_text:
            raise ValueError("No text could be extracted from the TXT file.")
            
        return full_text
    except Exception as e:
        logger.error(f"Error extracting TXT: {str(e)}")
        raise e

